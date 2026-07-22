#!/usr/bin/env python3
"""Generate the Potraces Subscription & Echo Limits .docx from the research JSON.
Usage: python make_docx.py data.json out.docx
Plain-English, founder-friendly. Tables for tiers / limits / screens / gaps.
"""
import json, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OLIVE = RGBColor(0x4F, 0x51, 0x04)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTE = RGBColor(0x6A, 0x6A, 0x6A)
HDRBG = "4F5104"
ZEBRA = "F3F3EE"

data = json.load(open(sys.argv[1], encoding="utf-8"))
out = sys.argv[2]

doc = Document()
# base style
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(11)
st.font.color.rgb = INK

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=10, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("" if text is None else str(text))
    r.bold = bold; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif color: r.font.color.rgb = color

def h1(t):
    p = doc.add_paragraph(); p.space_after = Pt(6)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(17); r.font.color.rgb = OLIVE
    return p

def h2(t):
    p = doc.add_heading(level=2)
    r = p.add_run(t); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = INK
    return p

def para(t, color=None, size=11, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.font.size = Pt(size); r.italic = italic; r.bold = bold
    if color: r.font.color.rgb = color
    return p

def bullet(t, color=None):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(t)
    if color: r.font.color.rgb = color
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; shade(c, HDRBG); set_cell(c, htext, bold=True, size=10, white=True)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val, size=10)
            if ri % 2 == 1: shade(cells[i], ZEBRA)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def src_line(sources):
    if sources:
        para("Source: " + "; ".join(sources[:6]), color=MUTE, size=8, italic=True)

# ---------- Cover ----------
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Potraces"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = OLIVE
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Subscription, Echo & Limits — Complete Guide"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = INK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("How the plans work, what each screen limits, what happens after the limit,\nand what we planned but haven't built yet."); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = MUTE
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Generated from the app's source code."); r.font.size = Pt(9); r.font.color.rgb = MUTE
doc.add_page_break()

# ---------- 1. Plans & prices ----------
h1("1.  The plans & prices")
tiers = data.get("tiers", {})
para(tiers.get("billingNotes","") or "The app sells a monthly or yearly subscription. Below are the tiers.", size=11)
rows = []
for t in tiers.get("tiers", []):
    rows.append([
        t.get("name",""), t.get("priceMonthly","—"), t.get("priceYearly","—"),
        t.get("yearlyPerMonth","—"), t.get("discount","—"), t.get("oneLineWhoItsFor","")
    ])
if rows:
    table(["Plan","Monthly","Yearly","Yearly / mo","Save","Who it's for"], rows,
          widths=[0.9,0.8,0.9,0.9,0.6,2.3])
src_line(tiers.get("sources"))

# ---------- 2. What is Echo ----------
h1("2.  What is “Echo”?")
echo = data.get("echo", {})
para(echo.get("whatIsEcho",""), size=11)
if echo.get("whereUsed"):
    h2("Where you can reach Echo")
    for w in echo["whereUsed"]: bullet(w)
h2("The free limit on Echo")
para("• Free chats: " + echo.get("freeChatLimit","—"))
para("• How a chat is counted: " + echo.get("howChatsAreCounted","—"))
para("• When it resets: " + echo.get("whenChatsReset","—"))
h2("What happens when you run out")
para(echo.get("afterLimit",""), bold=False)
if echo.get("smartAiByTier"):
    h2("Does the AI get smarter on higher plans?")
    para(echo["smartAiByTier"])
if echo.get("otherNotes"): para(echo["otherNotes"], color=MUTE, size=10, italic=True)
src_line(echo.get("sources"))

# ---------- 3. Every limit in the app ----------
h1("3.  Every limit, plan by plan")
limits = data.get("limits", {})
rows = []
for r_ in limits.get("resources", []):
    rows.append([
        r_.get("plainLabel",""), r_.get("free","—"), r_.get("basic","—"),
        r_.get("pro","—"), r_.get("premium","—"), r_.get("notes","")
    ])
if rows:
    table(["What's limited","Free","Basic","Pro","Premium","Notes"], rows,
          widths=[1.7,0.7,0.7,0.7,0.8,1.6])
if limits.get("resetPolicy"):
    para("Reset policy: " + limits["resetPolicy"], bold=True, size=10)
src_line(limits.get("sources"))

# ---------- 4. How the limit system works ----------
h1("4.  How the limit system works (behind the scenes)")
eng = data.get("engine", {})
para("How a limit is decided: " + eng.get("howLimitCheckWorks",""))
para("What happens at the limit: " + eng.get("whatHappensAtLimit",""))
if eng.get("hardVsSoftGate"): para("Hard block vs. soft nudge: " + eng["hardVsSoftGate"])
if eng.get("whenCountersReset"): para("When counters reset: " + eng["whenCountersReset"])
if eng.get("howIncremented"): para("How usage is counted up: " + eng["howIncremented"])
if eng.get("capabilityGatesVsQuotaGates"):
    h2("Two kinds of gate")
    para(eng["capabilityGatesVsQuotaGates"])
if eng.get("usageCountersTracked"):
    h2("Counters the app tracks")
    for c in eng["usageCountersTracked"]: bullet(c)
src_line(eng.get("sources"))

# ---------- 5. Screen by screen ----------
h1("5.  Screen by screen — what each one limits")
para("“Echo?” = does this screen use the AI. “After the limit” = what the user sees once they hit the cap.", color=MUTE, italic=True, size=10)
rows = []
for s in data.get("screens", []):
    rows.append([
        s.get("screen",""),
        "Yes" if s.get("usesEcho") else "—",
        (s.get("limitedResource","") if s.get("hasLimit") else "No limit"),
        (s.get("freeLimit","") if s.get("hasLimit") else "—"),
        s.get("afterLimitBehavior",""),
    ])
if rows:
    table(["Screen","Echo?","What's capped","Free limit","After the limit"], rows,
          widths=[1.7,0.6,1.4,0.9,2.0])

# detail cards for screens WITH limits
h2("The details, per screen")
for s in data.get("screens", []):
    if not s.get("hasLimit"): continue
    para(s.get("screen",""), bold=True, color=OLIVE, size=11)
    if s.get("limitedResource"): para("• Capped: " + s.get("limitedResource",""), size=10)
    if s.get("freeLimit"): para("• Free limit: " + s.get("freeLimit",""), size=10)
    if s.get("whenPaywallShows"): para("• Paywall appears when: " + s.get("whenPaywallShows",""), size=10)
    para("• After the limit: " + s.get("afterLimitBehavior",""), size=10)
    if s.get("hardOrSoft"): para("• Type: " + s.get("hardOrSoft",""), size=10)
    if s.get("notes"): para("  " + s["notes"], color=MUTE, italic=True, size=9)

# screens with NO limit
nolimit = [s for s in data.get("screens", []) if not s.get("hasLimit")]
if nolimit:
    h2("Screens with NO limit today")
    for s in nolimit:
        para("• " + s.get("screen","") + " — " + (s.get("notes","") or "no subscription gate."), size=10)

# ---------- 6. Planned but not built ----------
h1("6.  What we PLANNED to limit but haven't yet")
gaps = data.get("gaps", {})
if gaps.get("planSummary"): para(gaps["planSummary"], size=11)
rows = []
for g in gaps.get("gaps", []):
    rows.append([g.get("area",""), g.get("whatWasPlanned",""), g.get("whatExistsToday",""), g.get("status","")])
if rows:
    table(["Area","What we planned","What exists today","Status"], rows,
          widths=[1.4,2.1,2.1,0.9])
src_line(gaps.get("sources"))

# ---------- 7. Worth knowing ----------
h1("7.  Anything else worth knowing")
worth = data.get("worth", {})
for it in worth.get("items", []):
    para(it.get("topic",""), bold=True, size=11)
    para(it.get("detail",""), size=10)
src_line(worth.get("sources"))

doc.save(out)
print("WROTE", out)
