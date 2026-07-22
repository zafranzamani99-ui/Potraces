"""Generate docs/POTRACES_NOTIFICATIONS.docx — the full notification catalogue.
Temporary script: run once, then delete."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

ACCENT = RGBColor(0x3F, 0x5A, 0x1E)  # potraces-ish sage


def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    return h


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def kv_table(pairs):
    """2-col field/value table for one notification."""
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = 'Table Grid'
    t.autofit = False
    for i, (k, v) in enumerate(pairs):
        c0, c1 = t.rows[i].cells
        c0.width = Inches(1.7)
        c1.width = Inches(4.8)
        r = c0.paragraphs[0].add_run(k)
        r.bold = True
        c1.paragraphs[0].add_run(v)
    doc.add_paragraph()  # breathing room


def grid_table(headers, rows, widths=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'
    for j, htext in enumerate(headers):
        run = t.rows[0].cells[j].paragraphs[0].add_run(htext)
        run.bold = True
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            t.rows[i + 1].cells[j].paragraphs[0].add_run(cell)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()


# ───────────────────────── TITLE ─────────────────────────
title = doc.add_heading('Potraces Notifications — The Complete List', level=0)
p('Every notification Potraces can show or send: what fires it, when it fires, '
  'what it says, which screen it belongs to, and what happens when you tap it. '
  'Written from the actual code (client + server), 22 July 2026.', italic=True)

# ───────────────────────── BIG PICTURE ─────────────────────────
h1('1. The Big Picture')
p('Potraces can produce two kinds of notifications:')
bullets([
    'On-device (local) — the app schedules these on the phone itself. They fire at a set '
    'time even with no internet. Example: a bill reminder at 9:00 AM.',
    'Server push — the Potraces server (Supabase) sends these through the Expo push service '
    'the moment something happens. Example: a buyer places an order on your web shop link.',
])
p('Nothing on the server runs on a clock. There are no scheduled/cron push jobs — every '
  'push is triggered by a real event (an order placed, a payment confirmed, a button tapped). '
  'The only timed notifications are the local ones on the phone.')
p('On Android, server pushes arrive on named channels: "Pesanan" (orders) and "Kutipan" '
  '(Collectz), both high importance so they make sound and pop up. Channels named "Bill '
  'reminders", "Spending alerts" and "Payment reminders" also exist, but because of a small '
  'code quirk the local notifications never attach themselves to those channels, so on '
  'Android they actually arrive on the system default channel (details in section 6).')

# ───────────────────────── SUMMARY TABLE ─────────────────────────
h1('2. Everything At A Glance')
grid_table(
    ['#', 'Notification', 'Kind', 'When it fires', 'Where it lives'],
    [
        ['1', 'Bill due soon / overdue', 'Local', 'A few days before the bill date, at 9:00 AM', 'Bills & subscriptions'],
        ['2', 'Echo daily check-in', 'Local', 'Every day at your chosen time(s) — default 9:00 PM', 'Settings → Money Setup'],
        ['3', 'Spending above usual', 'Local', 'Checked each time the app opens (max once per day), fires immediately if a category is way above its 4-week average', 'Settings → Money Setup'],
        ['4', '"payment received?" QR nudge', 'Local', '10 minutes after a seller closes the QR sheet without confirming payment', 'Seller → Orders'],
        ['5', 'Quick Log test', 'Local', '1.5 seconds after pressing the test button', 'Quick Log setup → Diagnostics'],
        ['6', 'Pesanan Baru! (new web order)', 'Push', 'The instant a buyer submits your order page', 'Seller web shop link'],
        ['7', 'Payment received (QR)', 'Push', 'Seconds after a buyer pays your DuitNow QR online', 'Seller QR payments'],
        ['8', 'Quick Log confirmation', 'Push', 'About 1 second after you Quick-Log via the iPhone shortcut', 'Quick Log (Back Tap shortcut)'],
        ['9', 'Collectz — proof to review', 'Push', 'The instant a participant uploads payment proof', 'Collectz session'],
        ['10', 'Collectz — payment confirmed', 'Push', 'The moment the organiser confirms a payment', 'Collectz session'],
        ['11', 'Collectz — proof rejected', 'Push', 'The moment the organiser rejects a proof', 'Collectz session'],
        ['12', 'Collectz — pay reminder', 'Push', 'When the organiser taps Remind (max once per 24h)', 'Collectz session detail'],
        ['13', 'Collectz — updated / cancelled / settled', 'Push', 'When the organiser edits, cancels or settles the session', 'Collectz session'],
    ],
    widths=[0.3, 1.9, 0.7, 2.6, 1.5],
)

# ───────────────────────── PART 1: LOCAL ─────────────────────────
h1('3. On-Device Notifications (scheduled by the phone)')

h2('3.1 Bill & subscription reminders')
kv_table([
    ('What it is', 'One reminder per active recurring bill or subscription, fired a few days '
     'before the due date so it never sneaks up on you.'),
    ('When it fires', 'At 9:00 AM, N days before the next billing date. N defaults to 3 and can be '
     'changed per bill ("remind N days before" in the bill form). If that moment has already '
     'passed but the bill is due within 24 hours — or is already overdue — it fires once at the '
     'next 9:00 AM instead of staying silent. Paused bills and fully-paid instalment plans get none.'),
    ('What it says', 'Title: "Netflix is due soon" (or "…is overdue"). '
     'Body: "RM 45.00 due on 5 Jul 2026." (or "…was due 5 Jul 2026."). '
     'Uses your chosen currency. Always in English.'),
    ('Where it comes from', 'Automatic: armed at every app startup and re-synced 1 second after any '
     'change to your bills list. The app asks for notification permission at startup if you have bills.'),
    ('Tap it and…', 'The app just opens. There is no deep link into the bill itself.'),
    ('How to stop it', 'There is no on/off switch for bill reminders — pause or delete the bill, '
     'or deny notification permission at the OS level.'),
])

h2('3.2 Echo daily check-in')
kv_table([
    ('What it is', 'A gentle daily nudge from Echo to log your day\'s spending.'),
    ('When it fires', 'Every day at the time(s) you pick. Default is one reminder at 9:00 PM; '
     'you can add more times and remove them down to the last one. Repeats daily.'),
    ('What it says', 'Title: "Potraces". Body: "How did today go? Take 30 seconds to log your '
     'spending." Always in English, with the default notification sound.'),
    ('Where it comes from', 'Settings → Money Setup → "daily check-in" toggle (default OFF) and '
     '"remind me at" time chips. Turning it on with no times seeds 9:00 PM. If permission is '
     'missing you get a toast asking you to allow notifications in phone Settings.'),
    ('Tap it and…', 'Opens Echo chat in personal mode, where Echo greets you with a summary like '
     '"RM 45.00 across 3 today — anything before the day closes?" (that greeting is an in-chat '
     'message, not a notification).'),
    ('How to stop it', 'Turn the "daily check-in" toggle off.'),
])

h2('3.3 Spending alerts (category jumps above normal)')
kv_table([
    ('What it is', 'A heads-up that one of your spending categories is running unusually hot '
     'this week compared to your own recent history.'),
    ('When it fires', 'The check runs every time the app comes to the foreground, but at most once '
     'per 24 hours. A notification appears immediately if any category this week is above 150% of '
     'its trailing 4-week weekly average AND more than RM 20 above it. This is not a scheduled '
     'alarm — it only fires when you open the app and the numbers trip the threshold.'),
    ('What it says', 'Title: "spending above usual". Body: "this week\'s Food & Dining (+RM 85), '
     'Transport (+RM 42) is higher than your 4-week average." Up to 3 categories, biggest first. '
     'Always in English.'),
    ('Where it comes from', 'Settings → Money Setup → "Spending alerts" toggle (default ON). '
     'If notification permission was never granted it silently does nothing.'),
    ('Tap it and…', 'The app just opens. No deep link.'),
    ('How to stop it', 'Turn the "Spending alerts" toggle off.'),
])

h2('3.4 "payment received?" — QR nudge (seller)')
kv_table([
    ('What it is', 'A safety net for sellers using the manual QR standee: if you show a buyer your '
     'QR and close the sheet without confirming the money arrived, the app checks back with you.'),
    ('When it fires', 'Exactly 10 minutes after you close the QR sheet on an order that is still '
     'underpaid (static/manual QR flow only). One-off per order — if you confirm the payment, '
     'the reminder is cancelled automatically.'),
    ('What it says', 'Title: "payment received?" Body: "did the RM 25.00 payment for order #1042 '
     'arrive?" This one IS translated — in Malay: "bayaran diterima?" / "adakah bayaran RM 25.00 '
     'untuk pesanan #1042 sudah masuk?".'),
    ('Where it comes from', 'Nothing to set up — it is automatic in Seller → Orders when you use '
     'the manual QR flow.'),
    ('Tap it and…', 'Switches to business mode and opens that order in Seller → Orders.'),
    ('How to stop it', 'Confirm the payment (it cancels itself), or just ignore it — it fires once.'),
])

h2('3.5 Quick Log test notification')
kv_table([
    ('What it is', 'A manual self-test that proves two things: banners can appear on your phone, '
     'and tapping a Quick Log notification opens your transactions list.'),
    ('When it fires', '1.5 seconds after you tap "Send test notification".'),
    ('What it says', 'Title: "Potraces". Body: "Test — tap me to open your transactions!" '
     '(Malay: "Ujian — tekan saya untuk buka senarai transaksi!").'),
    ('Where it comes from', 'Quick Log setup screen → Diagnostics section → "Send test notification" '
     'button. Greyed out until notification permission is granted.'),
    ('Tap it and…', 'Opens the Transactions list in personal mode.'),
    ('How to stop it', 'Just don\'t press the button — it only ever fires on demand.'),
])

# ───────────────────────── PART 2: PUSH ─────────────────────────
h1('4. Server Push Notifications (sent when something happens)')
p('All of these travel: server → Expo push service → your phone. They need internet and a '
  'registered device token (the app registers your device silently at startup / setup). '
  'Dead tokens are cleaned up automatically by some senders.')

h2('4.1 "Pesanan Baru! 🛒" — new order from your web shop')
kv_table([
    ('What it is', 'Tells a seller the moment a buyer places an order through their public '
     'order page (the web shop link).'),
    ('When it fires', 'Instantly, the moment the order is saved. A database trigger on the '
     'orders table does it — there is no delay and nothing to arm.'),
    ('What it says', 'Title: "Pesanan Baru! 🛒". Body: "Ali baru letak pesanan RM 25.5" '
     '(customer name, then the total; "Pelanggan" if no name). This one is in Malay.'),
    ('Who gets it', 'The seller, on every device registered to their account.'),
    ('Tap it and…', 'Switches to business mode and opens that order in Seller → Orders. '
     'High-priority "Pesanan" Android channel with sound.'),
    ('How to stop it', 'There is no per-feature toggle; the master "Notifications" toggle only '
     'affects banners while the app is open. Stop sharing the order link, or sign out.'),
])

h2('4.2 "Payment received" — DuitNow QR paid (seller)')
kv_table([
    ('What it is', 'Confirms to a seller that an online QR payment actually landed, without '
     'them watching the screen.'),
    ('When it fires', 'Seconds after the payment provider (HitPay) confirms the buyer\'s DuitNow '
     'QR payment. The webhook first marks the order paid, then pushes. The provider retries '
     'automatically if the server hiccups.'),
    ('What it says', 'Title: "Payment received". Body: "MYR 12.50 — order #1042" (order number '
     'part omitted if there is none). In English.'),
    ('Who gets it', 'The seller, on all their registered devices.'),
    ('Tap it and…', 'Same as new orders: business mode → that order in Seller → Orders.'),
    ('How to stop it', 'No toggle. It only exists for orders paid through the online QR flow.'),
])

h2('4.3 Quick Log confirmation — "Logged RM4.50 out"')
kv_table([
    ('What it is', 'A receipt-in-your-pocket confirmation after you log an expense or income via '
     'the iPhone Quick Log shortcut (Back Tap), so you know it landed without opening the app.'),
    ('When it fires', 'About a second after the shortcut posts: the server saves the entry to your '
     'Quick Log inbox first, then pushes the confirmation.'),
    ('What it says', 'Title: "Logged RM4.50 out" (or "…in" for income). Body: "Kopi · 🍔 Food & '
     'Dining" — your note plus the category label. In English.'),
    ('Who gets it', 'You, on every device tied to your Quick Log key. Rate limit: 30 logs per '
     'minute per user.'),
    ('Tap it and…', 'Opens the Transactions list in personal mode (the entry itself is then '
     'drained from the inbox into your records).'),
    ('How to stop it', 'Remove the shortcut / revoke the Quick Log key in Quick Log setup.'),
])

h2('4.4 Collectz — "payment to review" (to the organiser)')
kv_table([
    ('What it is', 'Tells the organiser someone says they\'ve paid and their proof is waiting '
     'for review.'),
    ('When it fires', 'Instantly, when a participant uploads payment proof (their status flips '
     'to "pending"). Fired by a database trigger — no button needed.'),
    ('What it says', 'Title: "Collectz: payment to review". Body: "Ahmad uploaded proof for '
     '"Futsal Khamis"". In English.'),
    ('Who gets it', 'The session organiser, on all their devices.'),
    ('Tap it and…', 'Opens that session\'s detail screen (organiser view) in personal mode.'),
    ('How to stop it', 'No toggle — it is part of running a Collectz session.'),
])

h2('4.5 Collectz — "Payment confirmed ✅" (to a participant)')
kv_table([
    ('What it is', 'Tells a participant their payment was accepted.'),
    ('When it fires', 'The moment the organiser taps confirm on their pending payment.'),
    ('What it says', 'Title: "Payment confirmed ✅". Body: "Your payment for "Futsal Khamis" is '
     'confirmed". In English.'),
    ('Who gets it', 'That participant — only if they use the app (web-only participants can\'t '
     'be pushed). All their devices.'),
    ('Tap it and…', 'Opens the session on the Collectz join screen.'),
    ('How to stop it', 'No toggle.'),
])

h2('4.6 Collectz — "Payment rejected" (to a participant)')
kv_table([
    ('What it is', 'Tells a participant their proof was not accepted, so they can re-pay or '
     're-upload.'),
    ('When it fires', 'The moment the organiser rejects their proof.'),
    ('What it says', 'Title: "Payment rejected". Body: "Your proof for "Futsal Khamis" was '
     'rejected". In English.'),
    ('Who gets it', 'That participant (app users only), all devices.'),
    ('Tap it and…', 'Opens the session on the Collectz join screen.'),
    ('How to stop it', 'No toggle.'),
])

h2('4.7 Collectz — "Reminder: …" pay-up blast')
kv_table([
    ('What it is', 'The organiser\'s nag button: pings everyone who still owes money.'),
    ('When it fires', 'Only when the organiser taps the Remind button on the session detail '
     'screen. Cooldown: once per session per 24 hours — a second tap within 24h is refused.'),
    ('What it says', 'Title: "Reminder: Futsal Khamis". Body: "Your share is still unpaid — '
     'MYR 20.00. Tap to pay and upload proof." (amount omitted if no share can be computed). '
     'In English.'),
    ('Who gets it', 'Participants who are still unpaid or were rejected, have the app, and hold '
     'an active slot (not waiting list). Each gets it on all their devices.'),
    ('Tap it and…', 'Opens the session on the Collectz join screen, ready to pay/upload.'),
    ('How to stop it', 'Pay your share — or the organiser simply stops pressing the button.'),
])

h2('4.8 Collectz — session updated / cancelled / settled')
kv_table([
    ('What it is', 'Keeps every participant in sync when the organiser changes the plan, calls '
     'it off, or closes the books.'),
    ('When it fires', 'Updated: when the organiser saves an edit with the "notify changes" box '
     'ticked. Cancelled: when the organiser cancels the session (also sent just before a session '
     'is deleted, since the join link dies with it). Settled: when the organiser confirms the '
     'Settle dialog.'),
    ('What it says', 'Updated — title: "Updated: Futsal Khamis", body: the organiser\'s own '
     'message, or "The organizer updated this session — tap to see what changed." '
     'Cancelled — title: "Cancelled: Futsal Khamis", body: "The organizer cancelled this session." '
     'Settled — title: "Settled: Futsal Khamis", body: "All done — this session is now '
     'settled. 🎉". All in English.'),
    ('Who gets it', 'Every app-linked participant (any status), on all their devices.'),
    ('Tap it and…', 'Opens the session on the Collectz join screen.'),
    ('How to stop it', 'No toggle — they only fire on organiser actions.'),
])

# ───────────────────────── TAP ROUTING ─────────────────────────
h1('5. What Each Notification Opens When Tapped')
p('The app reads a small data tag inside every notification and routes you to the right screen. '
  'This works both when the app is open and when it was fully closed (cold start).')
grid_table(
    ['Notification tag', 'Tapping takes you to'],
    [
        ['echo_checkin (daily check-in)', 'Personal mode → Echo chat'],
        ['quick_log (Quick Log test + real confirmations)', 'Personal mode → Transactions list'],
        ['new_order / payment_received', 'Business mode → Seller → Orders, opened on that order'],
        ['collectz_pending (proof to review)', 'Personal mode → Collectz session detail (organiser view)'],
        ['collectz_confirmed / rejected / reminder / edited / cancelled / settled',
         'Personal mode → Collectz join screen for that session'],
        ['Bill reminders & spending alerts', 'No tag — the app simply opens to wherever you left it'],
    ],
    widths=[3.0, 3.5],
)

# ───────────────────────── SETTINGS ─────────────────────────
h1('6. The Switches That Control Notifications')
grid_table(
    ['Setting', 'Where', 'Default', 'What it actually does'],
    [
        ['Notifications (master)', 'Settings (shared app settings)', 'ON',
         'Affects banners while the app is open and the new-order toast. Does NOT stop scheduled '
         'local reminders or server pushes by itself.'],
        ['Daily check-in', 'Settings → Money Setup', 'OFF',
         'Arms/disarms the Echo daily reminder and holds your chosen time(s), default 9:00 PM.'],
        ['Spending alerts', 'Settings → Money Setup', 'ON',
         'Enables the once-a-day overspend check when the app opens.'],
        ['Bill reminder lead time', 'Each bill\'s form ("remind N days before")', '3 days',
         'Per-bill. There is no global on/off for bill reminders.'],
        ['OS permission', 'Phone Settings → Potraces', 'asked at startup / setup',
         'The real gate. Asked at app start (if you have bills), when enabling check-in, and in '
         'Quick Log setup.'],
    ],
    widths=[1.5, 1.7, 1.1, 2.2],
)

# ───────────────────────── GOOD TO KNOW ─────────────────────────
h1('7. Good To Know (quirks found in the code)')
bullets([
    'No red badge counts anywhere — the app never sets the app-icon badge.',
    'Languages are inconsistent: bill reminders, Echo check-in and spending alerts are always '
    'English; the QR "payment received?" nudge and the Quick Log test follow your app language '
    '(EN/MS); the new-order push is always Malay; every other push is always English.',
    'Android channels exist for bills / spending alerts / payment reminders, but the local '
    'notifications never attach themselves to a channel, so on Android they arrive on the system '
    'default channel. Server pushes DO use their channels ("Pesanan" for orders, "Kutipan" for '
    'Collectz).',
    'Two foreground banner handlers are registered and they disagree; in practice the "always '
    'show" one wins after startup, so banners appear while the app is open even if you turned '
    'the master Notifications toggle off. That is a code quirk, not intended behaviour.',
    'Bill reminders re-sync at every app start and 1 second after any bill change. Echo check-in '
    'reminders only re-sync when you change that setting. Spending alerts re-check at most once '
    'per 24 hours.',
    'Collectz has no automatic "event is starting soon" reminder — the only reminder is the '
    'organiser\'s manual Remind button (24h cooldown).',
    'Things that are NOT notifications: the "success/warning notification" calls in Goals and '
    'Notes are just phone vibrations (haptics); the debts "reminder" setting only shows in-app '
    'buttons on debt cards; Echo\'s check-in greeting inside the chat is a chat message, not a '
    'system notification.',
])

p('')
p('Source: full read of the Potraces client (src/services, src/screens, App.tsx) and the '
  'Supabase backend (edge functions + database triggers), 22 July 2026.', italic=True)

import os
out = os.path.join('docs', 'POTRACES_NOTIFICATIONS.docx')
doc.save(out)
print('saved', out)
