# -*- coding: utf-8 -*-
"""Generate Potraces-Bank-Connection-Plan.docx — a clear, visual brief for the
team/founder on what's possible, what's impossible, and what we're building re:
connecting wallets to banks / auto-capturing transactions.

Run:  python docs/research/generate_bank_connection_doc.py
Out:  docs/Potraces-Bank-Connection-Plan.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- On-brand palette (earthy, no harsh red) ----------------------------------
OLIVE        = "4F5104"   # accent / "possible now"
OLIVE_TINT   = "E9ECDA"
GOLD         = "B2780A"   # "possible but costly/partial"
GOLD_TINT    = "F6EBCE"
TERRA        = "C1694F"   # "blocked / impossible" — muted terracotta, NOT red
TERRA_TINT   = "F0E1DA"
BRONZE       = "8B7355"
INK          = "2A2A28"
MUTED        = "6B6A64"
WHITE        = "FFFFFF"
RULE         = "D8D6CE"

doc = Document()

# Base style
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10.5)
base.font.color.rgb = RGBColor.from_string(INK)


# ---- helpers ------------------------------------------------------------------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def cell_text(cell, text, bold=False, color=INK, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    for i, line in enumerate(str(text).split("\n")):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if align:
                p.alignment = align
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
    set_cell_margins(cell)


def para(text="", size=10.5, color=INK, bold=False, italic=False, after=6, before=0,
         align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullet(text, color=INK, size=10.5, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def heading(text, size=15, color=OLIVE, before=14, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def rule_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE)
    pbdr.append(bottom)
    pPr.append(pbdr)


def make_table(headers, rows, widths=None, header_fill=OLIVE, zebra=True,
               cell_size=9.5):
    """rows: list of [ (text, fill_or_None, text_color), ... ]"""
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        cell_text(hdr[i], h, bold=True, color=WHITE, size=cell_size)
        shade(hdr[i], header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, spec in enumerate(row):
            if isinstance(spec, tuple):
                txt, fill, tcol = (spec + (None, INK))[:3]
            else:
                txt, fill, tcol = spec, None, INK
            cell_text(cells[ci], txt, color=tcol or INK, size=cell_size)
            if fill:
                shade(cells[ci], fill)
            elif zebra and ri % 2 == 1:
                shade(cells[ci], "F7F6F2")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def legend():
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    items = [("POSSIBLE NOW", OLIVE_TINT, OLIVE),
             ("POSSIBLE — but costly / partial", GOLD_TINT, GOLD),
             ("IMPOSSIBLE until rails exist", TERRA_TINT, TERRA)]
    cells = t.rows[0].cells
    for i, (txt, fill, tcol) in enumerate(items):
        cell_text(cells[i], txt, bold=True, color=tcol, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(cells[i], fill)
    return t


# ================================ TITLE =======================================
p = para("POTRACES", size=12, color=BRONZE, bold=True, after=0)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title = para("Connecting Wallets to Banks", size=26, color=INK, bold=True, after=0)
para("What's possible, what's impossible, and what we're building",
     size=13, color=MUTED, italic=True, after=4)
para("Prepared June 2026  ·  Malaysia context  ·  Internal planning brief",
     size=9.5, color=MUTED, after=2)
rule_line()

# ================================ TL;DR =======================================
heading("The one-paragraph truth", size=14, before=10)
para(
    "Auto-capturing transactions (your wallet updating itself from the bank) is a "
    "REGULATORY capability, not an app trick. It only works where the government has "
    "built “open banking” rails. In Malaysia those rails arrive in 2027 for banks "
    "and ~2029 for e-wallets like TNG — they do not exist yet. So today, Potraces "
    "wallets are manual BY NECESSITY, not by oversight. There is nothing to connect to. "
    "What we CAN do now is make manual feel effortless (quick-add) and semi-automatic "
    "(reading bank emails) — and build the data model so that when the rails go live, "
    "real bank connection plugs in without a rewrite.",
    size=10.5, after=8)

legend()
para("", size=4, after=2)

# ===================== BIG PICTURE: POSSIBLE / IMPOSSIBLE ======================
heading("1.  At a glance — what is and isn't possible right now", size=14)
make_table(
    ["Capability", "Status today (June 2026)", "Why"],
    [
        [("Manual wallets (type a balance, log spending)", OLIVE_TINT, OLIVE),
         ("Working now", OLIVE_TINT, OLIVE),
         "This is what Potraces does today."],
        [("Faster manual capture (Back Tap / Siri / widget → “RM12 lunch”)", OLIVE_TINT, OLIVE),
         ("Buildable now", OLIVE_TINT, OLIVE),
         "iOS App Intents — no bank needed."],
        [("Read bank/e-wallet ALERT EMAILS → suggest a transaction", OLIVE_TINT, OLIVE),
         ("Buildable now", OLIVE_TINT, OLIVE),
         "Gmail/IMAP + your existing AI parser."],
        [("Apple Pay auto-capture (Tap-to-Pay)", GOLD_TINT, GOLD),
         ("Partial now", GOLD_TINT, GOLD),
         "iOS Shortcuts trigger; only Apple Pay spend."],
        [("Connect Maybank / CIMB bank accounts (live sync)", GOLD_TINT, GOLD),
         ("Possible, but costly", GOLD_TINT, GOLD),
         "Finverse/Brankas contract + compliance; 2 banks only."],
        [("Read TNG / GrabPay / other app NOTIFICATIONS on iPhone", TERRA_TINT, TERRA),
         ("Impossible", TERRA_TINT, TERRA),
         "iOS blocks reading other apps' notifications. No API."],
        [("Connect TNG e-wallet (live sync)", TERRA_TINT, TERRA),
         ("Impossible until ~2029", TERRA_TINT, TERRA),
         "E-wallets are last in BNM's timeline. No API exists."],
        [("Connect most other Malaysian banks (live sync)", TERRA_TINT, TERRA),
         ("Impossible until 2027/28", TERRA_TINT, TERRA),
         "PayNet's open-finance platform isn't live yet."],
    ],
    widths=[3.0, 1.7, 2.6],
)

# ===================== PAYMENT MIX MAPPING ====================================
heading("2.  Your real payment mix — can we auto-capture each one?", size=14)
para("Reflecting how Malaysians actually pay (Apple Pay rising, but TNG + bank QR dominate):",
     size=10, color=MUTED, italic=True, after=4)
make_table(
    ["Payment rail", "Best capture method", "iPhone", "Android", "Verdict"],
    [
        ["Apple Pay (growing)", "iOS Shortcuts ‘Transaction’ trigger",
         ("Auto", OLIVE_TINT, OLIVE), ("Auto (Google Pay)", OLIVE_TINT, OLIVE),
         ("Coverage grows over time", OLIVE_TINT, OLIVE)],
        ["Bank QR / DuitNow (MAE, CIMB)", "Open-banking aggregator + bank emails",
         ("Via aggregator", GOLD_TINT, GOLD), ("Yes", OLIVE_TINT, OLIVE),
         ("Best-covered long-term", OLIVE_TINT, OLIVE)],
        ["TNG eWallet (dominant)", "Manual / screenshot → AI / (Android push)",
         ("Manual only", TERRA_TINT, TERRA), ("Notif-read possible", GOLD_TINT, GOLD),
         ("The hard gap — ~2029", TERRA_TINT, TERRA)],
        ["Other banks", "Open Finance (PayNet) later",
         ("Wait 2027/28", TERRA_TINT, TERRA), ("Wait 2027/28", TERRA_TINT, TERRA),
         ("Rails not live yet", TERRA_TINT, TERRA)],
        ["Cash", "Manual quick-add",
         ("Manual", OLIVE_TINT, OLIVE), ("Manual", OLIVE_TINT, OLIVE),
         ("Always manual", OLIVE_TINT, OLIVE)],
    ],
    widths=[1.7, 2.3, 1.1, 1.2, 1.7],
    cell_size=8.8,
)
para("Key consequence: TNG — your single biggest rail — has no clean auto-capture on "
     "iPhone for years. The right answer for TNG is not an impossible ‘read’, it's "
     "FRICTIONLESS MANUAL CAPTURE. That's why the quick-add feature matters most.",
     size=10, color=BRONZE, bold=False, italic=True, before=4, after=4)

# ===================== WHY IMPOSSIBLE =========================================
heading("3.  Why some things are genuinely impossible (plain English)", size=14)
bullet("Apple's privacy sandbox: an iPhone app cannot read another app's notifications "
       "or SMS. There is no setting, no permission, no workaround. This is why the famous "
       "SMS-reading apps (e.g. Walnut) are Android-only.", bold_prefix="The iPhone wall — ")
bullet("Malaysia has no consumer bank-data API in production yet. Banks are not yet "
       "required to share your transaction data with apps. That requirement starts in 2027.",
       bold_prefix="No rails yet — ")
bullet("E-wallets (TNG, GrabPay, Boost) are scheduled LAST in the rollout (~2029) and "
       "currently expose no consumer data API and send push notifications (not emails) — "
       "the worst combination for capture on iPhone.", bold_prefix="E-wallets are last — ")

# ===================== MALAYSIA TIMELINE ======================================
heading("4.  The Malaysia timeline (BNM Open Finance)", size=14)
make_table(
    ["When", "Milestone", "What it unlocks"],
    [
        ["18 Nov 2025", "BNM Open Finance Exposure Draft published", "Rulebook drafted"],
        ["1 Mar 2026", "Industry consultation closes", "Rules finalised after"],
        ["Mid 2026", "PayNet platform pilot (7 banks + EPF)", "First live API tests"],
        [("1 Jan 2027", OLIVE_TINT, OLIVE), ("Banks with >1M customers go live", OLIVE_TINT, OLIVE),
         ("Connect big banks", OLIVE_TINT, OLIVE)],
        ["1 Jan 2028", "Banks with >100k customers", "More banks"],
        [("1 Jan 2029", GOLD_TINT, GOLD), ("E-wallets / EMIs (incl. TNG)", GOLD_TINT, GOLD),
         ("Connect TNG, finally", GOLD_TINT, GOLD)],
    ],
    widths=[1.2, 3.1, 2.0],
)
para("Standards confirmed: REST APIs, JSON, OAuth2, TLS 1.2+, ISO 20022 — i.e. exactly "
     "what an aggregator like Finverse already speaks, so integration later is standard work.",
     size=9.5, color=MUTED, italic=True, before=2)

# ===================== THE PLAN ===============================================
heading("5.  What we're building — the plan", size=14)
para("Four phases. The first two ship value now without any bank deal; the last two "
     "make us ‘open-finance-ready’ so 2027 is a plug-in, not a rebuild.",
     size=10, color=MUTED, italic=True, after=4)
make_table(
    ["Phase", "What", "Needs a bank deal?", "Effort", "Payoff"],
    [
        [("1", OLIVE_TINT, OLIVE), "App Intent quick-add (Back Tap / Siri / widget / Control Centre)",
         ("No", OLIVE_TINT, OLIVE), "Medium", "Instant manual capture; also receives Apple Pay"],
        [("2", OLIVE_TINT, OLIVE), "Email parsing (Gmail/IMAP → AI → suggested transaction)",
         ("No", OLIVE_TINT, OLIVE), "Medium", "Semi-auto across many banks + some e-wallets"],
        [("3", GOLD_TINT, GOLD), "Connection abstraction (future-ready wallet data model)",
         ("No", GOLD_TINT, GOLD), "Low", "Nothing breaks later; plug-in ready"],
        [("4", GOLD_TINT, GOLD), "Plug in open banking (Finverse now / PayNet 2027)",
         ("Yes", GOLD_TINT, GOLD), "High", "True live bank sync — when worth it"],
    ],
    widths=[0.6, 3.0, 1.2, 0.8, 2.0],
    cell_size=8.8,
)
para("Recommendation: build Phases 1–3 now. Hold Phase 4 until Open Finance is live "
     "and covers enough banks (and eventually TNG) to justify the recurring cost + compliance.",
     size=10, color=BRONZE, italic=True, before=4)

# ===================== CONNECTION ABSTRACTION =================================
heading("6.  The ‘Connection’ abstraction (how we stay future-ready)", size=14)
para("Today a wallet is a manual number. We add a few OPTIONAL fields so the same wallet "
     "can later be fed by a real bank connection — without changing its shape. No "
     "migration, nothing breaks for existing users.", size=10.5, after=4)
para("Wallet model — today vs. upgraded:", size=10.5, bold=True, after=3)
make_table(
    ["Field", "Today", "After upgrade"],
    [
        ["balance", "Manual number", "Manual OR fed by sync"],
        ["type", "Label only (bank/ewallet/credit)", "Same"],
        [("connectionId?", GOLD_TINT, GOLD), ("—", GOLD_TINT, GOLD), ("Links to a live data source", GOLD_TINT, GOLD)],
        [("externalAccountId?", GOLD_TINT, GOLD), ("—", GOLD_TINT, GOLD), ("The bank's account reference", GOLD_TINT, GOLD)],
        [("lastSyncedAt?", GOLD_TINT, GOLD), ("—", GOLD_TINT, GOLD), ("When it last auto-updated", GOLD_TINT, GOLD)],
        [("source?", GOLD_TINT, GOLD), ("—", GOLD_TINT, GOLD), ("manual / email / openbanking", GOLD_TINT, GOLD)],
    ],
    widths=[1.8, 2.6, 2.6],
)
para("The capture flow (same for every source):", size=10.5, bold=True, before=6, after=3)
para("source (email / Apple Pay / bank API)  →  normalise into one transaction shape  "
     "→  de-duplicate against what's already logged  →  show the user a suggestion  "
     "→  user taps Confirm  →  wallet balance updates.",
     size=10.5, after=4)
para("Golden rule: auto-captured items are SUGGESTIONS the user confirms — never silent "
     "writes. This keeps the app honest, calm, and trustworthy, and avoids double-counting.",
     size=10, color=BRONZE, italic=True)

# ===================== PHASE 1 & 2 DETAIL =====================================
heading("7.  Phase 1 — App Intent quick-add (build first)", size=13, color=OLIVE)
bullet("One small native module + Expo config plugin exposes an ‘Add expense’ action to iOS.")
bullet("That single action lights up Back Tap (double-tap back of iPhone), Siri (“hey Siri, "
       "add expense”), Spotlight, Control Centre, Lock Screen, and home-screen widgets.")
bullet("Same module RECEIVES Apple Pay transactions via a Shortcuts automation the user sets up once.")
bullet("Does double duty: the only good answer for TNG (manual) + the receiver for Apple Pay (auto).")
bullet("We provide a friendly one-time ‘set up auto-capture’ guide card.")

heading("8.  Phase 2 — Email parsing (build second)", size=13, color=OLIVE)
bullet("User connects Gmail (or any IMAP) or forwards alerts to a parse address.")
bullet("Reuse the existing AI receipt pipeline to read bank & e-wallet ALERT EMAILS and "
       "e-receipts → extract merchant, amount, date → suggest a transaction.")
bullet("Cross-platform (iPhone AND Android), no bank contract, no Apple sandbox problem.")
bullet("Covers many banks plus some e-wallet emails — the broadest cheap net available today.")
bullet("Privacy: parse on-device where possible, only financial emails, never personal mail, "
       "explicit consent, full BM/EN copy, calm framing, no red.")

# ===================== COST & EFFORT ==========================================
heading("9.  Honest cost & effort summary", size=14)
make_table(
    ["Option", "Up-front", "Ongoing cost", "Compliance", "Coverage"],
    [
        ["Quick-add (Phase 1)", "Native module work", ("None", OLIVE_TINT, OLIVE),
         ("None", OLIVE_TINT, OLIVE), "All rails (manual)"],
        ["Email parsing (Phase 2)", "Integration + AI prompts", ("AI tokens (small)", OLIVE_TINT, OLIVE),
         ("Consent / PDPA", GOLD_TINT, GOLD), "Many banks + some e-wallets"],
        ["Bank connect now (Finverse)", "OAuth + ingest + dedupe",
         ("Per-account fees", GOLD_TINT, GOLD), ("Data-consumer duties", TERRA_TINT, TERRA),
         ("Maybank + CIMB only", GOLD_TINT, GOLD)],
        ["Open Finance (2027+)", "Plug into Phase 3", ("TBD by PayNet", GOLD_TINT, GOLD),
         ("Registered access", GOLD_TINT, GOLD), ("Broad, growing", OLIVE_TINT, OLIVE)],
    ],
    widths=[1.9, 1.7, 1.5, 1.5, 1.7],
    cell_size=8.8,
)

# ===================== BOTTOM LINE ============================================
heading("10.  Bottom line", size=14)
bullet("Wallets are manual today because Malaysia has nothing to connect to yet — correct, "
       "not a flaw.", bold_prefix="Manual is correct — ", color=INK)
bullet("Build quick-add + email parsing now. They make the app feel automatic without a "
       "single bank deal, and they cover the TNG gap that bank connections can't.",
       bold_prefix="Win now — ", color=INK)
bullet("Add the few optional wallet fields now so 2027 is a plug-in. We become "
       "‘open-finance-ready from day one’ — the moat over SMS-scraping rivals.",
       bold_prefix="Future-proof — ", color=INK)
bullet("Don't sign an aggregator contract yet — two banks isn't worth the cost + "
       "compliance, and it still misses TNG. Revisit when Open Finance broadens coverage.",
       bold_prefix="Wait on bank connect — ", color=INK)

rule_line()
para("Full source research & citations: docs/research/auto-capture-global-research.md",
     size=8.5, color=MUTED, italic=True)
para("This brief is a plan, not a commitment of dates. The 2027/2029 dates are BNM's "
     "proposed timeline and may shift.", size=8.5, color=MUTED, italic=True)

# ---- save ---------------------------------------------------------------------
out_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "Potraces-Bank-Connection-Plan.docx")
doc.save(out_path)
print("WROTE:", out_path)
