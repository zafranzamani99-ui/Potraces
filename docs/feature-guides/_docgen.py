# -*- coding: utf-8 -*-
"""Shared Word-document toolkit for Potraces feature guides.

Every feature guide (add-expense, wallet, debt, ...) imports this so they share
one consistent, on-brand look. On-brand = earthy/calm palette, NO harsh red;
muted terracotta is used only as a subtle "manual / watch-out" tint.

Usage:
    import _docgen as dg
    d = dg.FeatureDoc(kicker="POTRACES · FEATURE GUIDE",
                      title="Wallets",
                      subtitle="Every way money is held and moved",
                      meta="Updated June 2026 · code-accurate reference")
    d.h("1.  What a wallet is")
    d.p("...")
    d.table([...], [...])
    d.save(r"c:\\path\\out.docx")
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- On-brand palette ---------------------------------------------------------
OLIVE      = "4F5104"
OLIVE_TINT = "E9ECDA"
GOLD       = "B2780A"
GOLD_TINT  = "F6EBCE"
TERRA      = "C1694F"   # muted terracotta — NEVER bright red
TERRA_TINT = "F0E1DA"
SKY        = "3F6E84"
SKY_TINT   = "DEE9EF"
BRONZE     = "8B7355"
INK        = "2A2A28"
MUTED      = "6B6A64"
WHITE      = "FFFFFF"
RULE       = "D8D6CE"
ZEBRA      = "F7F6F2"
CODEBG     = "F0EFEA"


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _cell_margins(cell, top=60, bottom=60, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _runs_into(p, text, size, color, bold=False, italic=False, mono=False):
    """Render text into paragraph p. Supports inline **bold** and `code`."""
    import re
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", str(text))
    for part in parts:
        if not part:
            continue
        b, it, code = bold, italic, mono
        t = part
        if part.startswith("**") and part.endswith("**"):
            b, t = True, part[2:-2]
        elif part.startswith("`") and part.endswith("`"):
            code, t = True, part[1:-1]
        r = p.add_run(t)
        r.bold = b
        r.italic = it
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        if code:
            r.font.name = "Consolas"
            r.font.size = Pt(size - 0.5)
            r.font.color.rgb = RGBColor.from_string(BRONZE)


class FeatureDoc:
    def __init__(self, kicker, title, subtitle, meta):
        doc = Document()
        base = doc.styles["Normal"]
        base.font.name = "Calibri"
        base.font.size = Pt(10.5)
        base.font.color.rgb = RGBColor.from_string(INK)
        # widen the page a touch
        sec = doc.sections[0]
        sec.left_margin = Inches(0.85)
        sec.right_margin = Inches(0.85)
        sec.top_margin = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        self.doc = doc

        self._line(kicker, 10, BRONZE, bold=True, after=1)
        self._line(title, 28, INK, bold=True, after=0)
        self._line(subtitle, 13.5, MUTED, italic=True, after=3)
        self._line(meta, 9, MUTED, after=2)
        self.rule()

    # ---- primitives ----------------------------------------------------------
    def _line(self, text, size, color, bold=False, italic=False, after=4, before=0):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        return p

    def h(self, text, color=OLIVE, size=15, before=15, after=5):
        return self._line(text, size, color, bold=True, before=before, after=after)

    def h2(self, text, color=BRONZE, size=12, before=9, after=3):
        return self._line(text, size, color, bold=True, before=before, after=after)

    def p(self, text="", size=10.5, color=INK, after=6, before=0, italic=False):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if text:
            _runs_into(p, text, size, color, italic=italic)
        return p

    def bullet(self, text, size=10.5, color=INK, after=3):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(after)
        _runs_into(p, text, size, color)
        return p

    def num(self, text, size=10.5, color=INK, after=3):
        p = self.doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(after)
        _runs_into(p, text, size, color)
        return p

    def rule(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), RULE)
        pbdr.append(bottom)
        pPr.append(pbdr)
        return p

    def callout(self, text, kind="note"):
        """kind: note(olive) | online(gold) | manual(terra) | tip(sky)"""
        fill, tcol, label = {
            "note":   (OLIVE_TINT, OLIVE,  "NOTE"),
            "online": (GOLD_TINT,  GOLD,   "ONLINE-ONLY"),
            "manual": (TERRA_TINT, TERRA,  "MANUAL / WATCH-OUT"),
            "tip":    (SKY_TINT,   SKY,    "TIP"),
        }[kind]
        t = self.doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = t.rows[0].cells[0]
        _shade(cell, fill)
        _cell_margins(cell, top=90, bottom=90, left=140, right=140)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label + "  ")
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(tcol)
        _runs_into(p, text, 9.8, INK)
        self.p("", after=2)
        return t

    def legend(self, items):
        """items: list of (text, fill, textcolor)"""
        t = self.doc.add_table(rows=1, cols=len(items))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        for i, (txt, fill, tcol) in enumerate(items):
            c = t.rows[0].cells[i]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt)
            r.bold = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor.from_string(tcol)
            _shade(c, fill)
            _cell_margins(c)
        self.p("", after=2)
        return t

    def table(self, headers, rows, widths=None, header_fill=OLIVE, zebra=True,
              size=9.3):
        """rows: list of rows; each cell is str OR (text, fill, textcolor)."""
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        t.autofit = False
        hdr = t.rows[0].cells
        for i, htext in enumerate(headers):
            c = hdr[i]
            c.text = ""
            pp = c.paragraphs[0]
            pp.paragraph_format.space_after = Pt(0)
            r = pp.add_run(htext)
            r.bold = True
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor.from_string(WHITE)
            _shade(c, header_fill)
            _cell_margins(c)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci, spec in enumerate(row):
                if isinstance(spec, tuple):
                    txt, fill, tcol = (spec + (None, INK))[:3]
                else:
                    txt, fill, tcol = spec, None, INK
                c = cells[ci]
                c.text = ""
                first = True
                for line in str(txt).split("\n"):
                    pp = c.paragraphs[0] if first else c.add_paragraph()
                    first = False
                    pp.paragraph_format.space_after = Pt(0)
                    pp.paragraph_format.space_before = Pt(0)
                    _runs_into(pp, line, size, tcol or INK)
                if fill:
                    _shade(c, fill)
                elif zebra and ri % 2 == 1:
                    _shade(c, ZEBRA)
                _cell_margins(c)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        self.p("", after=2)
        return t

    def save(self, path):
        self.doc.save(path)
        return path
